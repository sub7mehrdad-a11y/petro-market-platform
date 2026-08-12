"""
تبدیل یک فایل docx به یک لیست ترتیبی از بلوک‌ها (heading/paragraph/table)،
دقیقاً به همون ترتیبی که در سند اومده‌ن.

چرا لازم بود: python-docx به‌صورت پیش‌فرض پاراگراف‌ها و جدول‌ها رو در دو لیست
جدا می‌ده (document.paragraphs و document.tables) و ترتیب واقعی‌شون توی سند رو
گم می‌کنه. این ماژول مستقیم از روی XML سند به ترتیب واقعی می‌خونه.

تشخیص heading: یا از روی استایل Word («Heading 1»، «Heading 2»...)، یا — چون
بعضی گزارش‌ها (مثل گزارش عراق/اردن) اصلاً از استایل Heading استفاده نکرده‌ن،
فقط عنوان‌ها رو bold کرده‌ن — با این قاعده: پاراگرافی که کاملاً bold و کوتاه
(کمتر از ۱۲ کلمه) باشه هم heading در نظر گرفته می‌شه.
"""

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_block_items(document):
    """پاراگراف‌ها و جدول‌ها رو به ترتیب واقعی‌شون در سند yield می‌کنه."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _heading_level(paragraph: Paragraph):
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        digits = "".join(ch for ch in style_name if ch.isdigit())
        return int(digits) if digits else 1

    text = paragraph.text.strip()
    if not text:
        return None
    is_fully_bold = bool(paragraph.runs) and all(
        r.bold for r in paragraph.runs if r.text.strip()
    )
    word_count = len(text.split())
    if is_fully_bold and word_count <= 12:
        return 2
    return None


def extract_blocks(path: str) -> list[dict]:
    document = docx.Document(path)
    blocks = []

    for item in iter_block_items(document):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if not text:
                continue
            level = _heading_level(item)
            if level:
                blocks.append({"type": "heading", "level": level, "text": text})
            else:
                is_list = (item.style.name if item.style else "").startswith("List")
                blocks.append({"type": "list_item" if is_list else "paragraph", "text": text})
        elif isinstance(item, Table):
            headers = [c.text.strip() for c in item.rows[0].cells] if item.rows else []
            rows = [[c.text.strip() for c in row.cells] for row in item.rows[1:]]
            blocks.append({"type": "table", "headers": headers, "rows": rows})

    return blocks


def extract_plain_text(path: str) -> str:
    """برای دادن به مدل: فقط متن، بدون تفکیک نوع بلوک؛ جدول‌ها هم به شکل خط‌به‌خط میان."""
    blocks = extract_blocks(path)
    lines = []
    for b in blocks:
        if b["type"] == "table":
            lines.append(" | ".join(b["headers"]))
            for row in b["rows"]:
                lines.append(" | ".join(row))
        else:
            lines.append(b["text"])
    return "\n".join(lines)
