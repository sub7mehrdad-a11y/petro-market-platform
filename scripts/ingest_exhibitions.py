"""
ساخت بانک یکدست نمایشگاه‌ها در data/exhibitions.json از سه نوع منبع:

  1. جدول‌های تمیز (اکسل/جدول Word) — مثل نمایشگاه‌های کنیا و برزیل — مستقیم و
     بدون AI پارس می‌شن چون از قبل ساختاریافته‌ن.
  2. سند روایی («اطلس نمایشگاه‌های بین‌المللی») که هر نمایشگاه یک بلوک متنی جدا
     با عنوان بولد + پاراگراف‌های توضیحی + خط‌های کلیدواژه‌ای (تاریخ/محل/وب‌سایت)
     هست، نه جدول — این یکی چون ساختارش جدولی نیست، با Gemini استخراج می‌شه.

برای اضافه‌کردن کشور جدید با فایل جدولی (اکسل یا جدول Word): به TABLE_SOURCES
یک ورودی اضافه کن. برای سند روایی جدید: به NARRATIVE_SOURCES اضافه کن.
"""

import os
import re
import sys
import json

import openpyxl
import docx
from google import genai

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
SRC_DIR = os.path.join(BASE_DIR, "گزارش", "لیست نمایشگاه ها")
OUTPUT_FILE = os.path.join(BASE_DIR, "data", "exhibitions.json")

GEMINI_MODEL = "gemini-3.6-flash"

# --- منابع جدولی (بدون AI) ---
TABLE_SOURCES = [
    {
        "file": "نمایشگاه های کنیا.xlsx",
        "type": "xlsx",
        "country": "کنیا",
        "column_map": {
            "نام نمایشگاه": "name",
            "تاریخ برگزاری": "date",
            "مکان": "location",
            "تمرکز اصلی برای بازاریابی جوش شیرین": "focus",
        },
    },
    {
        "file": "Armenia_Exhibitions_SepehranChemical.xlsx",
        "type": "xlsx",
        "country": "ارمنستان",
        "header_row": 2,   # سطر ۱ عنوان فارسی، سطر ۲ خالی، سرستون‌ها در سطر ۳
        "column_map": {
            "نام نمایشگاه": "name",
            "حوزه تخصصی": "focus",
            "گرید هدف جوش شیرین": "target_grade",
            "خریداران و غرفه‌داران بالقوه": "audience",
            "تاریخ برگزاری بعدی": "date",
            "محل برگزاری": "location",
        },
    },
    {
        "file": "لیست نمایشگاه های کشور برزیل.docx",
        "type": "docx_table",
        "country": "برزیل",
        "column_map": {
            "نمایشگاه": "name",
            "برگزارکننده": "organizer",
            "شهر/محل": "location",
            "زمان معمول": "date",
            "grade هدف": "target_grade",
            "اهمیت": "focus",
        },
    },
    {
        "file": "Russia_Exhibitions.xlsx",
        "type": "xlsx",
        "country": "روسیه",
        "header_row": 1,   # سطر ۱ خالی، سرستون‌ها در سطر ۲
        "column_map": {
            "نام نمایشگاه": "name",
            "صنعت هدف (گرید مصرفی جوش شیرین)": "focus",
            "تاریخ برگزاری بعدی (میلادی)": "date",
            "وب‌سایت مرجع": "website",
        },
    },
]

# --- نمایشگاه‌هایی که در فایل‌های ورودی نبودن و با تحقیق وب اضافه شدن ---
# هر کدوم مستقیماً از سایت رسمی خود نمایشگاه تأیید شده (تاریخ و محل نقل مستقیم از
# همون صفحه). اگه نمایشگاه جدیدی پیدا کردید، همین‌جا اضافه کنید.
MANUAL_SOURCES = [
    {
        "name": "TURKCHEM Eurasia",
        "country": "ترکیه",
        "location": "استانبول، Istanbul Expo Center",
        "date": "۲۵ تا ۲۷ نوامبر ۲۰۲۶",
        "organizer": "Artkim Fuarcılık",
        "focus": "بزرگ‌ترین نمایشگاه صنعت شیمی ترکیه و اوراسیا؛ شامل مواد شیمیایی پایه، پتروشیمی و شیمی مواد غذایی — محل حضور مستقیم رقبای ترک و شناسایی توزیع‌کنندگان منطقه‌ای.",
        "website": "https://turkchem.com.tr/en",
        "source_file": "تحقیق وب (سایت رسمی نمایشگاه)",
    },
    {
        "name": "FNI — Food & Nutritional Ingredients",
        "country": "ترکیه",
        "location": "استانبول، Istanbul Expo Center (IFM)",
        "date": "۹ تا ۱۱ ژوئن ۲۰۲۷",
        "organizer": "Istanbul Expo Center",
        "focus": "نمایشگاه تخصصی مواد اولیه‌ی غذایی ترکیه — بازار هدف گرید خوراکی جوش شیرین (E500) و نقطه‌ی تماس با صنایع نانوایی و فرآوری غذایی منطقه.",
        "website": "https://fningredients.com/en",
        "source_file": "تحقیق وب (سایت رسمی نمایشگاه)",
    },
]

# --- منبع روایی (نیازمند Gemini) ---
NARRATIVE_SOURCES = [
    {
        "file": "گزارش تفصیلی اطلس نمایشگاه‌های بین‌المللی تخصصی (۲۰۲۶ - ۲۰۲۷).docx",
        "country": None,  # جهانیه؛ کشور هر نمایشگاه از دل متن استخراج می‌شه
    },
]

EXTRACTION_SYSTEM_PROMPT = """
تو یک استخراج‌کننده‌ی داده هستی. من متن یک سند رو می‌دم که هر نمایشگاه تجاری
رو در یک بلوک جدا توصیف کرده (عنوان + توضیح + اطلاعات تاریخ/محل/وب‌سایت).

برای هر نمایشگاهی که در متن پیدا کردی، یک آبجکت با این ساختار دقیق بساز:
{
  "name": "اسم نمایشگاه (بدون شهر/کشور توی پرانتز)",
  "country": "اسم کشور محل برگزاری (فقط از روی متن، حدس نزن)",
  "location": "شهر/محل برگزاری",
  "date": "تاریخ برگزاری همون‌طور که در متن اومده",
  "focus": "خلاصه‌ی یک تا دو جمله‌ای از تمرکز/اهمیت نمایشگاه برای بازاریابی جوش شیرین/سود پرک، به زبان خودت (نه کپی کامل پاراگراف)",
  "website": "وب‌سایت اگر ذکر شده، وگرنه null"
}

فقط یک JSON array از این آبجکت‌ها برگردون، بدون markdown fence و بدون توضیح اضافه.
چیزی رو از خودت اختراع نکن؛ اگه فیلدی توی متن نبود null بذار.
"""


def clean(v):
    if v is None:
        return None
    if isinstance(v, str):
        v = v.strip()
        return v or None
    return str(v).strip()


def load_xlsx_table(src):
    path = os.path.join(SRC_DIR, src["file"])
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.worksheets[0]
    rows = list(ws.iter_rows(values_only=True))

    # بعضی فایل‌ها (مثل نمایشگاه‌های ارمنستان) یک عنوان و یک سطر خالی بالای جدول
    # دارند؛ header_row (اندیس صفرمبنا) می‌گوید سرستون‌ها کجاست. پیش‌فرض سطر اول.
    header_idx = src.get("header_row", 0)
    header = [str(h).strip() if h else None for h in rows[header_idx]]

    records = []
    for row in rows[header_idx + 1:]:
        if not any(row):
            continue
        record = {"country": src["country"], "source_file": src["file"]}
        for idx, col_name in enumerate(header):
            field = src["column_map"].get(col_name)
            if field and idx < len(row):
                record[field] = clean(row[idx])
        if record.get("name"):
            records.append(record)
    return records


def load_docx_table(src):
    path = os.path.join(SRC_DIR, src["file"])
    d = docx.Document(path)
    if not d.tables:
        return []
    table = d.tables[0]
    header = [c.text.strip() for c in table.rows[0].cells]

    records = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        record = {"country": src["country"], "source_file": src["file"]}
        for idx, col_name in enumerate(header):
            field = src["column_map"].get(col_name)
            if field and idx < len(cells):
                record[field] = clean(cells[idx])
        if record.get("name"):
            records.append(record)
    return records


def extract_docx_text(path: str) -> str:
    d = docx.Document(path)
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())


def extract_json_array(text: str):
    fence_match = re.search(r"```(?:json)?\s*(\[.*?\])\s*```", text, re.DOTALL)
    raw = fence_match.group(1) if fence_match else text
    array_match = re.search(r"\[.*\]", raw, re.DOTALL)
    if not array_match:
        raise ValueError("پاسخ مدل شامل JSON array نبود:\n" + text[:500])
    return json.loads(array_match.group(0))


def load_narrative_exhibitions(src, client: genai.Client):
    path = os.path.join(SRC_DIR, src["file"])
    text = extract_docx_text(path)

    interaction = client.interactions.create(
        model=GEMINI_MODEL,
        system_instruction=EXTRACTION_SYSTEM_PROMPT,
        input=text,
    )
    records = extract_json_array(interaction.output_text)
    for r in records:
        r["source_file"] = src["file"]
    return records


def main():
    all_exhibitions = []

    for src in TABLE_SOURCES:
        loader = load_xlsx_table if src["type"] == "xlsx" else load_docx_table
        records = loader(src)
        print(f"[OK] {src['file']}: {len(records)} نمایشگاه")
        all_exhibitions.extend(records)

    if MANUAL_SOURCES:
        print(f"[OK] تحقیق وب: {len(MANUAL_SOURCES)} نمایشگاه")
        all_exhibitions.extend(MANUAL_SOURCES)

    if NARRATIVE_SOURCES:
        api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
        if not api_key:
            print("[WARN] GEMINI_API_KEY تنظیم نشده — منابع روایی رد می‌شن.")
        else:
            client = genai.Client(api_key=api_key)
            for src in NARRATIVE_SOURCES:
                records = load_narrative_exhibitions(src, client)
                print(f"[OK] {src['file']}: {len(records)} نمایشگاه (استخراج هوشمند)")
                all_exhibitions.extend(records)

    for i, e in enumerate(all_exhibitions, start=1):
        e["id"] = f"exhibition-{i}"

    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(all_exhibitions, f, ensure_ascii=False, indent=2)

    print(f"\n[DONE] مجموعاً {len(all_exhibitions)} نمایشگاه در {OUTPUT_FILE} ذخیره شد.")


if __name__ == "__main__":
    main()
